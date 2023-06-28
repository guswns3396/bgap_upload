import csv
from pathlib import Path
import re
import subprocess
import tempfile
from zipfile import BadZipFile

from . import RedcapUploader, RedcapUploaderError
import docx
<<<<<<< HEAD

__SOFFICE__ = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
=======
import platform

__SOFFICE__ = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
__SOFFICEWIN__ = r'C:\Program Files\LibreOffice\program\soffice'
>>>>>>> tmp


class QGlobalUploader(RedcapUploader):
    """Uploader for QGlobal .doc scoring files.

    This REDCap uploader handles the .doc scoring files produced by QGlobal,
    and implements the required functions of RedcapUploader.

    Members:
        _template_path, _api_url, _token, _log_path, and _uploaded_status are
        detailed in the constructor below.

        _score_doc is a docx.Document made from `score_path` in the constructor.
        _rc_map is a dictionary that maps REDCap variable names to positions
            in Word doc tables, identified by a tuple (table_index, column_name,
            row_name) where index is the 0-based index of the table in Word,
            and column_name and row_name are the strings in the target column
            and row.
    """

    def __init__(self, reports, template_path, api_url, token, log_path,
                 date_fields=None, uploaded_status=None, skip_complete=True):
        """Initializes a QGlobalToolboxUploader instance with data & API info.

        Arguments:
            template_path: Path to the appropriate template file.
            score_path: Path to the score file.
            api_url: API URL for the REDCap database, string.
            token: API token for the REDCap database, string.
            log_path: Path where we will log the JSON string that we send to
                REDCap for the push.
            uploaded_status: Status of a record after we have pushed it, see
                RedcapUploader class for constants, values are INCOMPLETE,
                UNVERIFIED, COMPLETE. Default value UNVERIFIED.
        """
        self._template_path = template_path
        self._reports = reports
        self._api_url = api_url
        self._token = token
        self._log_path = log_path
        if date_fields is None:
            self._date_fields = []
        else:
            self._date_fields = date_fields

        if not uploaded_status:
            self._uploaded_status = self.UNVERIFIED
        else:
            self._uploaded_status = uploaded_status
        self._skip_complete = skip_complete

        self._rc_map = self._parse_template()

        super().__init__()

    def pull(self):
        errors = []
        pulled_data = []

        for (subj_id, event), report in self._reports.items():
            print(f'Pulling {subj_id}, {event}, {report.report_path}')
            try:
                data = self._parse_score_file(report.report_path,
                                              subj_id, event)
                print(f'Parsing for {subj_id}, {event} successful')
            except RedcapUploaderError as err:
                errors.append(err)
                continue

            if data:
                data[self.id_field()] = subj_id
                data[self.event_field()] = event
                pulled_data.append(data)

        return pulled_data, errors

    def _parse_template(self):
        """Parses a QGlobal template to create a variable -> table cell mapping.

        This function expects a csv file with a header row of:
            rc_var,table_idx,row_name,col_name
        where table_idx is the 0-based index of the table in Word (indexed by
        order of appearance in the document), and row_name and col_name are the
        strings in the target cell's column and row. The value in the cell
        identified by that table index, row name, and column name will be mapped
        to the REDCap variable rc_var.

        Returns:
            A dictionary that maps REDCap variable names to the appropriate
            table index, column, and row, identified by a tuple (table_index,
            column_name, row_name).
        """
        with open(self._template_path) as template_file:
            rc_map = {
                    r['rc_var']: {'table_idx': int(r['table_idx']),
                                  'row_name': r['row_name'],
                                  'col_name': r['col_name'],
                                  'header_offset': int(r['header_offset']),
                                  'optional_field': r['optional_field'] != '0',
                                  'convert_age': r['convert_age'] != '0'}
                    for r in csv.DictReader(template_file)
            }

        return rc_map

    def _parse_score_file(self, score_path, subj_id, event):
        try:
            score_doc = docx.Document(score_path)
        except (ValueError, BadZipFile):
            if score_path.exists():
                print(f'Converting {score_path} to docx...')
                with tempfile.TemporaryDirectory() as tmpdir:
                    docx_path = convert_doc2docx(score_path, Path(tmpdir))
                    score_doc = docx.Document(docx_path)
            else:
                raise QGlobalUploaderError('Scoring file not found.',
                                           form_path=score_path)

        rc_vals = {}

        print(f'Parsing {subj_id}, {event}')

        for rc_var, rc_var_info in self._rc_map.items():
            if ( self._skip_complete and
                 self.is_complete(subj_id, event, rc_var) ):
                continue

            table_idx = rc_var_info['table_idx']
            row_name = rc_var_info['row_name'].replace('\\n', '\n')
            col_name = rc_var_info['col_name'].replace('\\n', '\n')
            table = score_doc.tables[table_idx]
            header_row = table.rows[rc_var_info['header_offset']].cells

            # Special case: looking for a value inside a cell
            if row_name.startswith(': '):
                col_idx = [idx for idx, cell in enumerate(header_row)
                           if col_name in cell.text]
            else:
                col_idx = [idx for idx, cell in enumerate(header_row)
                           if cell.text == col_name]

            if len(col_idx) == 0:
                raise QGlobalUploaderError(
                    f'No column {col_name} in table {table_idx}.',
                    subj_id=subj_id, event=event, form_path=score_path
                )
            else:
                # Use the first one even if there are multiple
                col_idx = col_idx[0]

            # Find the row and value
            # Special case: looking for a value inside a cell
            if row_name.startswith(': '):
                anchor_re = re.compile(f'[\n\t]?{row_name[2:]}:(.+)[\t\n]?')
                for row in table.rows:
                    anchor_match = anchor_re.search(row.cells[col_idx].text)
                    if anchor_match:
                        rc_vals[rc_var] = anchor_match.group(1).strip()
                        if rc_var_info['convert_age']:
                            rc_vals[rc_var] = convert_age(rc_vals[rc_var])
                        break
                else:
                    raise QGlobalUploaderError(
                            f'{row_name[2:]} not found in {col_name}.',
                            subj_id=subj_id, event=event, form_path=score_path)

            # Standard case
            else:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    if row_text[0] == row_name:
                        rc_vals[rc_var] = row_text[col_idx].strip()
                        if rc_var_info['convert_age']:
                            rc_vals[rc_var] = convert_age(rc_vals[rc_var])
                        break
                else:
                    if not rc_var_info['optional_field']:
                        raise QGlobalUploaderError(
                            f'No row {row_name} in table {table_idx}.',
                            subj_id=subj_id, event=event, form_path=score_path
                        )

            completed_field = self.completed_field(rc_var)
            rc_vals[completed_field] = self._uploaded_status

        return rc_vals

    def api_url(self):
        return self._api_url

    def token(self):
        return self._token

    def log_path(self):
        return self._log_path

    def date_fields(self):
        return self._date_fields

    def change_log_path(self, new_log_path):
        self._log_path = new_log_path


def convert_age(age):
    age_match = re.fullmatch(r'(<)?(\d+):(\d+)(\+)?', age)
    if not age_match:
        raise ValueError(f'{age} does not match the pattern <years>:<months>.')

    (lessthan, years, months, plus) = age_match.groups('')

    if int(months) >= 12:
        raise ValueError(
                f'Months should be in the range of 0 to 11 (got {months}).')

    return f'{lessthan}{int(years) + (int(months) / 12):.4f}{plus}'


def convert_doc2docx(in_path, out_dir):
    # if not windows
    if platform.system() != 'Windows':
        process = subprocess.run(
                [__SOFFICE__, '--convert-to', 'docx', '--outdir', out_dir, in_path],
                 capture_output=True)
    # if windows
    else:
        process = subprocess.run(
            [__SOFFICEWIN__, '--convert-to', 'docx', '--outdir', out_dir, in_path],
            capture_output=True)
    if process.returncode != 0:
        raise ValueError(
                f'soffice failed to convert {in_path} to a docx file: '
                f'\n\nstdout: {process.stdout}\n\nstderr: {process.stderr}\n')
    else:
        docx_path = out_dir / in_path.with_suffix('.docx').name
        assert(docx_path.exists())
        return docx_path


class QGlobalReport:
    def __init__(self, report_path):
        self.report_path = report_path


class QGlobalUploaderError(RedcapUploaderError):
    pass


def test():
    data_path = Path('/Users/picc/workspace/bgap/data')
    redcap_path = Path('/Users/picc/Projects/KSTRT/Data/REDCapUploads')
    token = (redcap_path / 'token.txt').read_text().rstrip('\r\n')
    api_url = 'https://redcap.stanford.edu/api/'

    for report in ('BASC3PRS', 'BASC3SRP', 'Vineland3'):
        score_reports = [{'path': data_path / f'{report}_Report_12345_1.doc',
                          'subj_id': '12345',
                          'event': 'year_1_arm_1'},
                        ]
        template_path = data_path / 'templates' / f'{report}_Template.csv'
        log_path = data_path / 'logs' / f'test_log_{report}.txt'

        if report == 'Vineland3':
            date_fields = ('vineland_date', )
        elif report == 'BASC3PRS':
            date_fields = ('basc_prs_date', )
        elif report == 'BASC3SRP':
            date_fields = ('basc_srp_date', )

        qg = QGlobalToolboxUploader(
                template_path, score_reports, api_url, token, log_path,
                date_fields=date_fields)
        sent, resp, errs = qg.push()
        print(sent)
        print(resp)
        print(errs)
